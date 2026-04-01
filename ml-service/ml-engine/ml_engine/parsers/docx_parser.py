"""
DOCX parser using python-docx.
"""

import logging
from pathlib import Path
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base_parser import BaseParser
from ml_engine.utils import ResumeParserError

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):

    def parse(self, file_path: str | Path) -> str:
        try:
            path = self._validate_file(file_path)
            doc = Document(str(path))   
            
            text_lines = []

            for paragraph in doc.paragraphs:
                content = paragraph.text.strip()
                if content:
                    text_lines.append(content)

            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(
                        cell.text.strip() 
                        for cell in row.cells 
                        if cell.text and cell.text.strip()
                    )
                    if row_text:
                        text_lines.append(row_text)
            # safety guard for extremely large documents
            if len(text_lines) > 5000:
                logger.warning(f"DOCX content too large to safely process: {file_path}")
                raise ResumeParserError("DOCX content too large to safely process")

            if len(text_lines) < 3:
                logger.warning(f"DOCX parsing returned sparse content: {file_path}")
                raise ResumeParserError("DOCX parsing returned empty content")

            text = "\n".join(text_lines)
            text = text.replace("\r", "\n")

            # normalize blank lines
            text = "\n".join(line for line in text.splitlines())

            return text

        except PackageNotFoundError as exc:
            msg = f"Invalid or corrupted DOCX file: {file_path}"
            logger.error(msg)
            raise ResumeParserError(msg) from exc

        except ResumeParserError:
            raise

        except Exception as exc:
            msg = f"Unexpected DOCX parsing failure for file: {file_path}"
            logger.exception(msg)
            raise ResumeParserError(msg) from exc